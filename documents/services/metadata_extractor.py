# documents/services/metadata_extractor.py
import os
import magic
import hashlib
from datetime import datetime
from pathlib import Path
import logging

# Imports pour différents types de fichiers
try:
    import PyPDF2
    from PyPDF2 import PdfReader
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import openpyxl
    from openpyxl import load_workbook
except ImportError:
    openpyxl = None

try:
    from PIL import Image
    from PIL.ExifTags import TAGS
except ImportError:
    Image = None
    TAGS = None

logger = logging.getLogger('documents')


class MetadataExtractor:
    """Service d'extraction de métadonnées des documents"""

    def __init__(self):
        self.supported_types = {
            'application/pdf': self._extract_pdf_metadata,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx_metadata,
            'application/msword': self._extract_doc_metadata,
            'text/plain': self._extract_text_metadata,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._extract_xlsx_metadata,
            'application/vnd.ms-excel': self._extract_xls_metadata,
            'image/jpeg': self._extract_image_metadata,
            'image/png': self._extract_image_metadata,
        }

    def extract_metadata(self, file_path):
        """
        Extrait les métadonnées d'un fichier

        Args:
            file_path (str): Chemin vers le fichier

        Returns:
            dict: Métadonnées extraites
        """
        try:
            # Métadonnées de base
            metadata = self._get_basic_metadata(file_path)

            # Détection du type MIME
            mime_type = magic.from_file(file_path, mime=True)
            metadata['mime_type'] = mime_type

            # Extraction spécifique selon le type
            if mime_type in self.supported_types:
                specific_metadata = self.supported_types[mime_type](file_path)
                metadata.update(specific_metadata)
            else:
                logger.warning(f"Type de fichier non supporté: {mime_type}")
                metadata['warning'] = f"Type de fichier non supporté: {mime_type}"

            logger.info(f"Métadonnées extraites pour: {file_path}")
            return metadata

        except Exception as e:
            logger.error(f"Erreur lors de l'extraction des métadonnées: {str(e)}")
            return {
                'error': str(e),
                'extracted_at': datetime.now().isoformat()
            }

    def extract_full_content(self, file_path, max_chars=None):
        """
        Extrait le contenu textuel complet d'un fichier SANS LIMITATION

        Args:
            file_path (str): Chemin vers le fichier
            max_chars (int): Limite de caractères (ignorée - pour compatibilité)

        Returns:
            str: Contenu textuel complet INTÉGRAL
        """
        try:
            mime_type = magic.from_file(file_path, mime=True)
            content = ""

            if mime_type == 'application/pdf':
                content = self._extract_full_pdf_content(file_path)
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                content = self._extract_full_docx_content(file_path)
            elif mime_type == 'text/plain':
                content = self._extract_full_text_content(file_path)
            elif mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                content = self._extract_full_xlsx_content(file_path)
            else:
                logger.warning(f"Extraction complète non supportée pour: {mime_type}")
                content = ""

            # SUPPRESSION DE TOUTE LIMITATION - EXTRACTION COMPLÈTE
            logger.info(f"Contenu extrait: {len(content)} caractères au total")
            return content

        except Exception as e:
            logger.error(f"Erreur extraction contenu complet: {str(e)}")
            return ""

    def _extract_full_pdf_content(self, file_path):
        """Extrait tout le contenu textuel d'un PDF"""
        if not PyPDF2:
            return ""

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                full_text = []

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            full_text.append(f"--- Page {page_num + 1} ---\n{text}")
                    except Exception as e:
                        logger.warning(f"Erreur extraction page {page_num + 1}: {e}")
                        continue

                return "\n\n".join(full_text)

        except Exception as e:
            logger.error(f"Erreur extraction PDF complète: {str(e)}")
            return ""

    def _extract_full_docx_content(self, file_path):
        """Extrait tout le contenu textuel d'un fichier DOCX"""
        if not DocxDocument:
            return ""

        try:
            doc = DocxDocument(file_path)
            full_text = []

            # Extraction des paragraphes
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)

            # Extraction du contenu des tableaux
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))

                if table_text:
                    full_text.append("--- Tableau ---\n" + "\n".join(table_text))

            return "\n\n".join(full_text)

        except Exception as e:
            logger.error(f"Erreur extraction DOCX complète: {str(e)}")
            return ""

    def _extract_full_text_content(self, file_path):
        """Extrait tout le contenu d'un fichier texte"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Essayer avec d'autres encodages
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except:
                    continue
            return ""
        except Exception as e:
            logger.error(f"Erreur extraction texte complet: {str(e)}")
            return ""

    def _extract_full_xlsx_content(self, file_path):
        """Extrait le contenu textuel d'un fichier Excel"""
        if not openpyxl:
            return ""

        try:
            workbook = load_workbook(file_path, read_only=True, data_only=True)
            full_text = []

            for sheet_name in workbook.sheetnames:
                try:
                    sheet = workbook[sheet_name]
                    sheet_text = [f"--- Feuille: {sheet_name} ---"]

                    for row in sheet.iter_rows(values_only=True):
                        row_text = []
                        for cell_value in row:
                            if cell_value is not None and str(cell_value).strip():
                                row_text.append(str(cell_value).strip())

                        if row_text:
                            sheet_text.append(" | ".join(row_text))

                    if len(sheet_text) > 1:  # Plus que juste le titre
                        full_text.append("\n".join(sheet_text))

                except Exception as e:
                    logger.warning(f"Erreur extraction feuille {sheet_name}: {e}")
                    continue

            workbook.close()
            return "\n\n".join(full_text)

        except Exception as e:
            logger.error(f"Erreur extraction XLSX complète: {str(e)}")
            return ""

    def _get_basic_metadata(self, file_path):
        """Extrait les métadonnées de base du fichier"""
        file_stat = os.stat(file_path)
        file_path_obj = Path(file_path)

        # Calcul du hash MD5
        md5_hash = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                md5_hash.update(chunk)

        return {
            'filename': file_path_obj.name,
            'file_extension': file_path_obj.suffix.lower(),
            'file_size': file_stat.st_size,
            'created_at': datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
            'md5_hash': md5_hash.hexdigest(),
            'extracted_at': datetime.now().isoformat(),
        }

    def _extract_pdf_metadata(self, file_path):
        """Extrait les métadonnées d'un fichier PDF"""
        if not PyPDF2:
            return {'error': 'PyPDF2 non installé'}

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)

                metadata = {
                    'document_type': 'PDF',
                    'num_pages': len(pdf_reader.pages),
                    'is_encrypted': pdf_reader.is_encrypted,
                }

                # Métadonnées du document
                if pdf_reader.metadata:
                    pdf_info = pdf_reader.metadata
                    metadata.update({
                        'title': pdf_info.get('/Title', ''),
                        'author': pdf_info.get('/Author', ''),
                        'subject': pdf_info.get('/Subject', ''),
                        'creator': pdf_info.get('/Creator', ''),
                        'producer': pdf_info.get('/Producer', ''),
                        'creation_date': str(pdf_info.get('/CreationDate', '')),
                        'modification_date': str(pdf_info.get('/ModDate', '')),
                    })

                # Extraction du texte de la première page pour analyse
                if len(pdf_reader.pages) > 0:
                    first_page = pdf_reader.pages[0]
                    text_preview = first_page.extract_text()[:500]
                    metadata['text_preview'] = text_preview
                    metadata['estimated_word_count'] = len(text_preview.split()) * len(pdf_reader.pages)

                return metadata

        except Exception as e:
            logger.error(f"Erreur extraction PDF: {str(e)}")
            return {'error': f'Erreur extraction PDF: {str(e)}'}

    def _extract_docx_metadata(self, file_path):
        """Extrait les métadonnées d'un fichier DOCX"""
        if not DocxDocument:
            return {'error': 'python-docx non installé'}

        try:
            doc = DocxDocument(file_path)

            metadata = {
                'document_type': 'DOCX',
                'num_paragraphs': len(doc.paragraphs),
                'num_tables': len(doc.tables),
            }

            # Propriétés du document
            core_props = doc.core_properties
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'comments': core_props.comments or '',
                'keywords': core_props.keywords or '',
                'category': core_props.category or '',
                'created': core_props.created.isoformat() if core_props.created else '',
                'modified': core_props.modified.isoformat() if core_props.modified else '',
            })

            # Extraction du texte pour analyse
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)

            text_content = '\n'.join(full_text)
            metadata['text_preview'] = text_content[:500]
            metadata['word_count'] = len(text_content.split())
            metadata['character_count'] = len(text_content)

            return metadata

        except Exception as e:
            logger.error(f"Erreur extraction DOCX: {str(e)}")
            return {'error': f'Erreur extraction DOCX: {str(e)}'}

    def _extract_doc_metadata(self, file_path):
        """Extrait les métadonnées d'un fichier DOC (ancien format Word)"""
        # Pour les fichiers .doc, on fait une extraction basique
        return {
            'document_type': 'DOC',
            'note': 'Extraction limitée pour les fichiers .doc',
        }

    def _extract_text_metadata(self, file_path):
        """Extrait les métadonnées d'un fichier texte"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

            lines = content.split('\n')
            words = content.split()

            return {
                'document_type': 'TEXT',
                'line_count': len(lines),
                'word_count': len(words),
                'character_count': len(content),
                'text_preview': content[:500],
                'encoding': 'utf-8',
            }

        except UnicodeDecodeError:
            # Essayer avec d'autres encodages
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()

                    lines = content.split('\n')
                    words = content.split()

                    return {
                        'document_type': 'TEXT',
                        'line_count': len(lines),
                        'word_count': len(words),
                        'character_count': len(content),
                        'text_preview': content[:500],
                        'encoding': encoding,
                    }
                except:
                    continue

            return {'error': 'Impossible de décoder le fichier texte'}

        except Exception as e:
            logger.error(f"Erreur extraction TEXT: {str(e)}")
            return {'error': f'Erreur extraction TEXT: {str(e)}'}

    def _extract_xlsx_metadata(self, file_path):
        """Extrait les métadonnées d'un fichier Excel XLSX"""
        if not openpyxl:
            return {'error': 'openpyxl non installé'}

        try:
            workbook = load_workbook(file_path, read_only=True)

            metadata = {
                'document_type': 'XLSX',
                'num_worksheets': len(workbook.worksheets),
                'worksheet_names': [ws.title for ws in workbook.worksheets],
            }

            # Propriétés du document
            props = workbook.properties
            metadata.update({
                'title': props.title or '',
                'creator': props.creator or '',
                'description': props.description or '',
                'subject': props.subject or '',
                'keywords': props.keywords or '',
                'category': props.category or '',
                'created': props.created.isoformat() if props.created else '',
                'modified': props.modified.isoformat() if props.modified else '',
            })

            # Analyse de la première feuille
            if workbook.worksheets:
                first_sheet = workbook.worksheets[0]
                metadata.update({
                    'first_sheet_name': first_sheet.title,
                    'max_row': first_sheet.max_row,
                    'max_column': first_sheet.max_column,
                })

            workbook.close()
            return metadata

        except Exception as e:
            logger.error(f"Erreur extraction XLSX: {str(e)}")
            return {'error': f'Erreur extraction XLSX: {str(e)}'}

    def _extract_xls_metadata(self, file_path):
        """Extrait les métadonnées d'un fichier Excel XLS"""
        return {
            'document_type': 'XLS',
            'note': 'Extraction limitée pour les fichiers .xls',
        }

    def _extract_image_metadata(self, file_path):
        """Extrait les métadonnées d'une image"""
        if not Image:
            return {'error': 'Pillow non installé'}

        try:
            with Image.open(file_path) as img:
                metadata = {
                    'document_type': 'IMAGE',
                    'format': img.format,
                    'mode': img.mode,
                    'size': img.size,
                    'width': img.width,
                    'height': img.height,
                }

                # Extraction des données EXIF
                if hasattr(img, '_getexif') and img._getexif():
                    exif_data = img._getexif()
                    if exif_data:
                        exif_metadata = {}
                        for tag_id, value in exif_data.items():
                            tag = TAGS.get(tag_id, tag_id)
                            exif_metadata[tag] = str(value)
                        metadata['exif'] = exif_metadata

                return metadata

        except Exception as e:
            logger.error(f"Erreur extraction IMAGE: {str(e)}")
            return {'error': f'Erreur extraction IMAGE: {str(e)}'}